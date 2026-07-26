import os
import tempfile
import pytest
import docx
from unittest.mock import AsyncMock, patch, MagicMock

from models.resume import (
    ParsedResumeData,
    EducationItem,
    ExperienceItem,
    ProjectItem,
    CertificationItem,
    AchievementItem,
)
from services.resume_service import ResumeService


def test_fallback_regex_parse():
    sample_text = """
    John Doe
    Email: john.doe@example.com
    Phone: (123) 456-7890
    LinkedIn: https://www.linkedin.com/in/johndoe
    GitHub: https://github.com/johndoe
    Summary: Experienced Python and React software engineer proficient in SQL, Docker, and AWS.
    """
    data = ResumeService.fallback_regex_parse(sample_text)
    assert data.email == "john.doe@example.com"
    assert data.phone == "(123) 456-7890"
    assert data.linkedin is not None and "linkedin.com/in/johndoe" in data.linkedin
    assert data.github is not None and "github.com/johndoe" in data.github
    assert "Python" in data.skills
    assert "React" in data.skills
    assert "Docker" in data.skills
    assert "Aws" in data.skills


def test_extract_text_docx():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        doc = docx.Document()
        doc.add_paragraph("Jane Doe")
        doc.add_paragraph("Senior Backend Developer using FastAPI and MongoDB.")
        doc.save(tmp_path)

        text = ResumeService.extract_text_from_file(tmp_path)
        assert "Jane Doe" in text
        assert "FastAPI and MongoDB" in text
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_unsupported_file_format():
    with pytest.raises(ValueError, match="Unsupported file format"):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            ResumeService.extract_text_from_file(tmp_path)
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)


def test_pydantic_schema_validation():
    sample_json = {
        "name": "Alice Smith",
        "email": "alice@example.com",
        "phone": "555-0199",
        "linkedin": "linkedin.com/in/alicesmith",
        "github": "github.com/alicesmith",
        "education": [
            {
                "degree": "B.S. in Computer Science",
                "institution": "MIT",
                "start_date": "2018",
                "end_date": "2022",
                "description": "GPA 3.9"
            }
        ],
        "experience": [
            {
                "role": "Software Engineer",
                "company": "Tech Corp",
                "start_date": "2022",
                "end_date": "Present",
                "description": "Developed backend APIs",
                "technologies": ["Python", "FastAPI"]
            }
        ],
        "skills": ["Python", "FastAPI", "MongoDB"],
        "projects": [
            {
                "title": "Digital Twin App",
                "description": "AI resume and career app",
                "technologies": ["React", "FastAPI", "Gemini"],
                "link": "github.com/alicesmith/digitaltwin"
            }
        ],
        "certifications": [
            {
                "name": "AWS Certified Developer",
                "issuer": "Amazon",
                "date": "2023"
            }
        ],
        "achievements": [
            {
                "title": "Hackathon Winner",
                "description": "First place in AI track"
            }
        ]
    }

    parsed = ParsedResumeData.model_validate(sample_json)
    assert parsed.name == "Alice Smith"
    assert len(parsed.education) == 1
    assert parsed.education[0].institution == "MIT"
    assert len(parsed.experience) == 1
    assert parsed.experience[0].role == "Software Engineer"
    assert len(parsed.projects) == 1
    assert parsed.projects[0].title == "Digital Twin App"
    assert len(parsed.certifications) == 1
    assert parsed.certifications[0].name == "AWS Certified Developer"
    assert len(parsed.achievements) == 1
    assert parsed.achievements[0].title == "Hackathon Winner"


@pytest.mark.asyncio
async def test_analyze_resume_ai_fallback():
    sample_text = "John Doe john.doe@example.com Python Javascript"
    with patch("services.resume_service.get_llm_provider") as mock_provider:
        # Simulate an exception in LLM generation
        mock_instance = MagicMock()
        mock_instance.generate_text = AsyncMock(side_effect=Exception("API Quota Exceeded"))
        mock_provider.return_value = mock_instance

        parsed_data, method = await ResumeService.analyze_resume_with_ai(sample_text)
        assert method == "regex_fallback"
        assert parsed_data.email == "john.doe@example.com"
        assert "Python" in parsed_data.skills
        assert "Javascript" in parsed_data.skills


@pytest.mark.asyncio
async def test_analyze_resume_ai_success():
    sample_text = "Alice Smith alice@example.com Python"
    fake_llm_json = """```json
    {
      "name": "Alice Smith",
      "email": "alice@example.com",
      "phone": "555-0100",
      "linkedin": null,
      "github": null,
      "education": [],
      "experience": [],
      "skills": ["Python", "Docker"],
      "projects": [],
      "certifications": [],
      "achievements": []
    }
    ```"""
    with patch("services.resume_service.get_llm_provider") as mock_provider:
        mock_instance = MagicMock()
        mock_instance.generate_text = AsyncMock(return_value=fake_llm_json)
        mock_provider.return_value = mock_instance

        parsed_data, method = await ResumeService.analyze_resume_with_ai(sample_text)
        assert method == "ai"
        assert parsed_data.name == "Alice Smith"
        assert parsed_data.email == "alice@example.com"
        assert "Python" in parsed_data.skills
        assert "Docker" in parsed_data.skills


@pytest.mark.asyncio
async def test_process_resume():
    sample_text = "Alice Smith alice@example.com Python React"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        doc = docx.Document()
        doc.add_paragraph(sample_text)
        doc.save(tmp_path)

        mock_db = MagicMock()
        mock_db.profiles.find_one = AsyncMock(return_value={"user_id": "user123", "skills": ["Git"]})
        mock_db.profiles.update_one = AsyncMock()
        mock_db.resumes.insert_one = AsyncMock(return_value=MagicMock(inserted_id="resume123"))

        response, record = await ResumeService.process_resume("user123", tmp_path, mock_db)
        assert response.message == "Resume uploaded and analyzed successfully"
        assert response.file_url == tmp_path
        assert "Git" not in response.extracted_skills # response only contains extracted skills
        assert "Python" in response.extracted_skills
        assert "React" in response.extracted_skills

        # Verify db calls
        mock_db.profiles.update_one.assert_called_once()
        mock_db.resumes.insert_one.assert_called_once()
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
